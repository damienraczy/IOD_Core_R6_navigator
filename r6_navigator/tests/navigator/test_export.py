import pytest
from pathlib import Path

from docx import Document as DocxDocument

from r6_navigator.navigator.services import crud
from r6_navigator.navigator.services.export_docx import (
    ExportConfig,
    export_bulk,
    export_capacity,
    make_filename,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _docx_text(path: Path) -> str:
    """Returns all paragraph text joined by newlines (includes headings)."""
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _docx_headings(path: Path, level: int = 1) -> list[str]:
    doc = DocxDocument(str(path))
    style_name = f"Heading {level}"
    return [p.text for p in doc.paragraphs if p.style.name == style_name]


def _docx_paragraphs_by_style(path: Path, style: str) -> list[str]:
    doc = DocxDocument(str(path))
    return [p.text for p in doc.paragraphs if p.style.name == style]


# ---------------------------------------------------------------------------
# Fixture: session with I1a capacity (FR translation with non-empty definition)
# ---------------------------------------------------------------------------

@pytest.fixture
def cap_session(session):
    crud.create_capacity(
        session, "I", 1, "a",
        label="Démontrer la responsabilité",
        lang="fr",
        definition="- Premier point.\n- Deuxième point.",
        central_function="Fonction centrale de test",
    )
    return session


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_export_single_capacity_fr(cap_session, tmp_path):
    config = ExportConfig(output_path=tmp_path / "test_fr.docx")
    path = export_capacity("I1a", cap_session, config)

    assert path.exists()
    assert path.stat().st_size > 0


def test_export_single_capacity_en(cap_session, tmp_path):
    crud.upsert_capacity_translation(
        cap_session, "I1a", "en",
        label="Demonstrating Responsibility",
        definition="- English point one.\n- English point two.",
    )
    config = ExportConfig(output_path=tmp_path / "test_en.docx", language="en")
    path = export_capacity("I1a", cap_session, config)

    assert path.exists()
    assert path.stat().st_size > 0


def test_export_includes_fiche_content(cap_session, tmp_path):
    config = ExportConfig(output_path=tmp_path / "test_fiche.docx")
    path = export_capacity("I1a", cap_session, config)

    assert "Premier point." in _docx_text(path)


def test_export_bullet_rendering(cap_session, tmp_path):
    """Definition stored as '- phrase.\n' must produce multiple List Bullet paragraphs."""
    config = ExportConfig(output_path=tmp_path / "test_bullets.docx")
    path = export_capacity("I1a", cap_session, config)

    bullets = _docx_paragraphs_by_style(path, "List Bullet")
    assert len(bullets) >= 2, "Definition should produce at least 2 bullet paragraphs"
    assert "Premier point." in bullets
    assert "Deuxième point." in bullets


def test_export_includes_questions(cap_session, tmp_path):
    crud.create_question(cap_session, "I1a", text="Question de test", lang="fr")

    config = ExportConfig(output_path=tmp_path / "test_questions.docx")
    path = export_capacity("I1a", cap_session, config)

    assert "Question de test" in _docx_text(path)


def test_export_observable_items_under_questions(cap_session, tmp_path):
    """Observable items must appear in the document when include_questions=True."""
    crud.create_observable_item(cap_session, "I1a", "OK", "Comportement OK", lang="fr")

    # Items présents si include_questions=True
    config = ExportConfig(output_path=tmp_path / "with_questions.docx")
    path = export_capacity("I1a", cap_session, config)
    assert "Comportement OK" in _docx_text(path)

    # Items absents si include_questions=False
    config_no_q = ExportConfig(
        output_path=tmp_path / "no_questions.docx",
        include_questions=False,
    )
    path_no_q = export_capacity("I1a", cap_session, config_no_q)
    assert "Comportement OK" not in _docx_text(path_no_q)


def test_export_includes_coaching(cap_session, tmp_path):
    crud.upsert_coaching_translation(
        cap_session, "I1a", "fr", reflection_themes="- Thème de test."
    )

    config = ExportConfig(output_path=tmp_path / "test_coaching.docx")
    path = export_capacity("I1a", cap_session, config)

    assert "Thème de test." in _docx_text(path)


def test_export_excludes_fiche_when_flagged(cap_session, tmp_path):
    config = ExportConfig(
        output_path=tmp_path / "test_no_fiche.docx",
        include_fiche=False,
    )
    path = export_capacity("I1a", cap_session, config)

    assert "Premier point." not in _docx_text(path)


def test_export_bulk(session, tmp_path):
    """export_bulk génère un fichier par capacité dans le répertoire cible."""
    crud.create_capacity(session, "I", 1, "a", label="Cap I1a", lang="fr")
    crud.create_capacity(session, "O", 1, "a", label="Cap O1a", lang="fr")

    output_dir = tmp_path / "bulk"
    config = ExportConfig(output_path=output_dir)
    paths = export_bulk(["I1a", "O1a"], session, config)

    assert len(paths) == 2
    for p in paths:
        assert p.exists()
        assert p.stat().st_size > 0


def test_export_both_languages(cap_session, tmp_path):
    """language='both' génère un fichier FR et un fichier EN pour chaque capacité."""
    crud.upsert_capacity_translation(
        cap_session, "I1a", "en",
        label="Demonstrating Responsibility",
        definition="- English definition.",
    )

    output_dir = tmp_path / "both"
    config = ExportConfig(output_path=output_dir, language="both")
    paths = export_bulk(["I1a"], cap_session, config)

    assert len(paths) == 2
    langs_in_names = {p.name.rsplit("_", 1)[-1].replace(".docx", "") for p in paths}
    assert langs_in_names == {"fr", "en"}


def test_export_empty_fields_no_crash(session, tmp_path):
    """No EN translation exists — all fields empty — must not raise."""
    crud.create_capacity(session, "I", 1, "a", label="FR only", lang="fr")

    config = ExportConfig(
        output_path=tmp_path / "test_empty.docx",
        language="en",
    )
    path = export_capacity("I1a", session, config)

    assert path.exists()


def test_make_filename():
    """make_filename construit un nom de fichier sanitisé."""
    name = make_filename("I1a", "Démontrer la responsabilité", "fr")
    assert name.startswith("I1a_")
    assert name.endswith("_fr.docx")
    assert " " not in name
