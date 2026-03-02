"""Service d'export DOCX pour R6 Navigator.

Génère un fichier Word par capacité avec les sections Fiche, Questions et Coaching.
Le nom de fichier est auto-construit depuis l'identifiant, le libellé et la langue.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from r6_navigator.navigator.services import crud


# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

@dataclass
class ExportConfig:
    """Configuration de l'export DOCX.

    Args:
        output_path: Chemin de destination — fichier pour ``export_capacity``,
            répertoire pour ``export_bulk``.
        language: Langue du document exporté : ``"fr"``, ``"en"`` ou ``"both"``.
            Avec ``"both"``, ``export_bulk`` génère un fichier par langue pour
            chaque capacité ; ``export_capacity`` inclut les deux langues dans
            un seul document.
        include_fiche: Inclure la section Fiche (définition, fonction centrale,
            observable, risques).
        include_questions: Inclure les manifestations observables et les questions STAR.
        include_coaching: Inclure la section Coaching (thèmes, leviers, missions).
    """

    output_path: Path
    language: str = "fr"
    include_fiche: bool = True
    include_questions: bool = True
    include_coaching: bool = True


# ---------------------------------------------------------------------------
# Localised labels for the DOCX document body
# ---------------------------------------------------------------------------

_LABELS: dict[str, dict[str, str]] = {
    "fr": {
        "level": "Niveau",
        "axis": "Axe",
        "pole": "Pôle",
        "definition": "Définition",
        "central_function": "Fonction centrale",
        "observable_items": "Manifestations observables",
        "risk_insufficient": "Risque si insuffisant",
        "risk_excessive": "Risque si excessif",
        "questions": "Questions",
        "reflection_themes": "Thèmes de réflexion",
        "intervention_levers": "Leviers d'intervention",
        "recommended_missions": "Missions à envisager",
    },
    "en": {
        "level": "Level",
        "axis": "Axis",
        "pole": "Pole",
        "definition": "Definition",
        "central_function": "Central function",
        "observable_items": "Observable manifestations",
        "risk_insufficient": "Risk if insufficient",
        "risk_excessive": "Risk if excessive",
        "questions": "Questions",
        "reflection_themes": "Reflection themes",
        "intervention_levers": "Intervention levers",
        "recommended_missions": "Recommended missions",
    },
}

_CATEGORY_LABELS: dict[str, dict[str, str]] = {
    "fr": {
        "OK": "Comportements efficaces",
        "EXC": "Risque si excessif",
        "DEP": "Dépasse les attentes",
        "INS": "Risque si insuffisant",
    },
    "en": {
        "OK": "Effective behaviors",
        "EXC": "Excessive risk",
        "DEP": "Above expectations",
        "INS": "Insufficient risk",
    },
}

# Ordre d'affichage des catégories dans le document exporté.
_CATEGORY_ORDER = ("OK", "DEP", "EXC", "INS")


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------

def make_filename(capacity_id: str, label: str, lang: str) -> str:
    """Construit un nom de fichier DOCX depuis l'identifiant, le libellé et la langue.

    Args:
        capacity_id: Identifiant de la capacité (ex. ``"I1a"``).
        label: Libellé de la capacité (peut contenir des caractères spéciaux).
        lang: Code de la langue (``"fr"`` ou ``"en"``).

    Returns:
        Nom de fichier de la forme ``"I1a_Libellé_fr.docx"``, avec le libellé
        tronqué à 40 caractères et les espaces remplacés par des tirets bas.
    """
    safe = re.sub(r"[^\w\s-]", "", label).strip()
    safe = re.sub(r"\s+", "_", safe)[:40]
    return f"{capacity_id}_{safe}_{lang}.docx"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def export_capacity(capacity_id: str, session: Session, config: ExportConfig) -> Path:
    """Génère un fichier DOCX pour une seule capacité.

    Si ``config.language == "both"``, le document contient le contenu en français
    suivi du contenu en anglais.

    Args:
        capacity_id: Identifiant de la capacité à exporter (ex. ``"I1a"``).
        session: Session SQLAlchemy active.
        config: Configuration de l'export (langue, sections, chemin).

    Returns:
        Chemin absolu du fichier DOCX généré.
    """
    doc = Document()
    if config.language == "both":
        _add_capacity_to_doc(doc, capacity_id, session, config, lang="fr")
        _add_capacity_to_doc(doc, capacity_id, session, config, lang="en")
    else:
        _add_capacity_to_doc(doc, capacity_id, session, config)
    doc.save(str(config.output_path))
    return config.output_path


def export_bulk(
    capacity_ids: list[str], session: Session, config: ExportConfig
) -> list[Path]:
    """Génère un fichier DOCX par capacité (et par langue si ``"both"``) dans un répertoire.

    Le répertoire ``config.output_path`` est créé s'il n'existe pas.
    Le nom de chaque fichier est construit par ``make_filename()``.

    Args:
        capacity_ids: Liste des identifiants de capacités à exporter.
        session: Session SQLAlchemy active.
        config: Configuration de l'export ; ``output_path`` doit être un répertoire.

    Returns:
        Liste des chemins absolus des fichiers DOCX générés.
    """
    output_dir = config.output_path
    output_dir.mkdir(parents=True, exist_ok=True)

    langs = ["fr", "en"] if config.language == "both" else [config.language]
    paths: list[Path] = []

    for capacity_id in capacity_ids:
        for lang in langs:
            trans = crud.get_capacity_translation(session, capacity_id, lang)
            label = trans.label if trans and trans.label else capacity_id
            filename = make_filename(capacity_id, label, lang)
            doc = Document()
            _add_capacity_to_doc(doc, capacity_id, session, config, lang=lang)
            dest = output_dir / filename
            doc.save(str(dest))
            paths.append(dest)

    return paths


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _add_bullet_text(doc: Document, value: str) -> None:
    """Convertit une chaîne au format ``"- phrase.\\n"`` en paragraphes Word bullets.

    Chaque ligne commençant par ``-`` devient un paragraphe de style ``List Bullet``.
    Les lignes vides ou ne commençant pas par ``-`` sont ignorées.

    Args:
        doc: Document Word en cours de construction.
        value: Texte en format liste à puces (``"- item1.\\n- item2.\\n"``).
    """
    for line in value.splitlines():
        text = line.strip()
        if text.startswith("- "):
            text = text[2:].strip()
        elif text.startswith("-"):
            text = text[1:].strip()
        else:
            continue  # Ignore les lignes sans tiret.
        if text:
            doc.add_paragraph(text, style="List Bullet")


def _add_capacity_to_doc(
    doc: Document,
    capacity_id: str,
    session: Session,
    config: ExportConfig,
    lang: str | None = None,
) -> None:
    """Ajoute le contenu d'une capacité au document Word.

    Args:
        doc: Document Word en cours de construction.
        capacity_id: Identifiant de la capacité à ajouter.
        session: Session SQLAlchemy active.
        config: Configuration de l'export (sections incluses).
        lang: Langue à utiliser pour cette passe ; si None, utilise ``config.language``.
    """
    if lang is None:
        lang = config.language

    labels = _LABELS.get(lang, _LABELS["fr"])
    cat_labels = _CATEGORY_LABELS.get(lang, _CATEGORY_LABELS["fr"])

    capacity = crud.get_capacity(session, capacity_id)
    if capacity is None:
        return

    trans = crud.get_capacity_translation(session, capacity_id, lang)
    label = trans.label if trans and trans.label else capacity_id

    # ── Heading 1 : identifiant — libellé ─────────────────────────────────────
    doc.add_heading(f"{capacity_id} — {label}", level=1)

    # ── Tableau de métadonnées (Niveau / Axe / Pôle) ──────────────────────────
    meta = doc.add_table(rows=3, cols=2)
    try:
        meta.style = "Table Grid"
    except KeyError:
        pass
    meta.cell(0, 0).text = labels["level"]
    meta.cell(0, 1).text = capacity.level_code
    meta.cell(1, 0).text = labels["axis"]
    meta.cell(1, 1).text = str(capacity.axis_number)
    meta.cell(2, 0).text = labels["pole"]
    meta.cell(2, 1).text = capacity.pole_code

    # ── Section Fiche ──────────────────────────────────────────────────────────
    if config.include_fiche and trans:
        # Définition (liste à puces).
        if (trans.definition or "").strip():
            doc.add_heading(labels["definition"], level=2)
            _add_bullet_text(doc, trans.definition)

        # Fonction centrale (texte libre, pas de puces).
        if (trans.central_function or "").strip():
            doc.add_heading(labels["central_function"], level=2)
            doc.add_paragraph(trans.central_function)

        # Risques (listes à puces).
        for field_name, label_key in [
            ("risk_insufficient", "risk_insufficient"),
            ("risk_excessive", "risk_excessive"),
        ]:
            value = getattr(trans, field_name, None) or ""
            if value.strip():
                doc.add_heading(labels[label_key], level=2)
                _add_bullet_text(doc, value)

    # ── Section Questions ──────────────────────────────────────────────────────
    if config.include_questions:
        # Manifestations observables groupées par catégorie (Heading 2 + Heading 3).
        items_all = crud.get_observable_items(session, capacity_id)
        if items_all:
            doc.add_heading(labels["observable_items"], level=2)
            for code in _CATEGORY_ORDER:
                items_cat = [i for i in items_all if i.category_code == code]
                if not items_cat:
                    continue
                cat_title = f"{cat_labels.get(code, code)}"
                doc.add_heading(cat_title, level=3)
                for item in items_cat:
                    item_trans = crud.get_observable_item_translation(
                        session, item.item_id, lang
                    )
                    text = item_trans.text if item_trans else ""
                    if text.strip():
                        doc.add_paragraph(text, style="List Bullet")

        # Questions STAR (liste numérotée).
        questions = crud.get_questions(session, capacity_id)
        if questions:
            doc.add_heading(labels["questions"], level=2)
            for idx, question in enumerate(questions, start=1):
                q_trans = crud.get_question_translation(
                    session, question.question_id, lang
                )
                text = q_trans.text if q_trans else ""
                doc.add_paragraph(f"{idx}. {text}")

    # ── Section Coaching ───────────────────────────────────────────────────────
    if config.include_coaching:
        coaching_trans = crud.get_coaching_translation(session, capacity_id, lang)
        if coaching_trans:
            for field_name, label_key in [
                ("reflection_themes", "reflection_themes"),
                ("intervention_levers", "intervention_levers"),
                ("recommended_missions", "recommended_missions"),
            ]:
                value = getattr(coaching_trans, field_name, None) or ""
                if value.strip():
                    doc.add_heading(labels[label_key], level=2)
                    _add_bullet_text(doc, value)
