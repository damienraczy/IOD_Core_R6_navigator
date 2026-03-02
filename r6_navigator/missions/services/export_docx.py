"""Service d'export DOCX pour le module Missions.

Génère un fichier Word de rapport de mission avec page de garde,
tableau de profil S-O-I, narratif et annexe des extraits analysés.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


_MISSION_LABELS: dict[str, dict[str, str]] = {
    "fr": {
        "cover_title": "Rapport de mission R6",
        "client": "Client",
        "date": "Date",
        "profile_title": "Profil de maturité S-O-I",
        "capacity": "Capacité",
        "avg_score": "Score moyen",
        "nb_obs": "Nb observations",
        "directions": "Directions",
        "narrative_title": "Analyse narrative",
        "extracts_title": "Annexe — Extraits analysés",
        "extract": "Extrait",
        "validation": "Statut",
    },
    "en": {
        "cover_title": "R6 Mission Report",
        "client": "Client",
        "date": "Date",
        "profile_title": "S-O-I Maturity Profile",
        "capacity": "Capacity",
        "avg_score": "Avg score",
        "nb_obs": "Nb observations",
        "directions": "Directions",
        "narrative_title": "Narrative analysis",
        "extracts_title": "Annex — Analyzed extracts",
        "extract": "Extract",
        "validation": "Status",
    },
}


def export_mission_report_docx(
    mission_id: str,
    session_factory,
    output_path: Path,
    lang: str = "fr",
) -> Path:
    """Génère un DOCX de rapport de mission.

    Structure : page de garde → tableau profil S-O-I → narratif → annexe extraits.

    Args:
        mission_id: ID de la mission.
        session_factory: Callable retournant une session SQLAlchemy.
        output_path: Chemin de destination du fichier DOCX.
        lang: Langue du document (``"fr"`` ou ``"en"``).

    Returns:
        Chemin absolu du fichier créé.
    """
    from collections import defaultdict
    from datetime import date

    from r6_navigator.missions.services import crud

    labels = _MISSION_LABELS.get(lang, _MISSION_LABELS["fr"])
    doc = Document()

    with session_factory() as session:
        mission = crud.get_mission(session, mission_id)
        if mission is None:
            raise ValueError(f"Mission '{mission_id}' not found")

        client_name = mission.client_name
        report = crud.get_latest_report(session, mission_id)
        report_content = report.content if report else ""

        interpretations = crud.list_interpretations_for_mission(session, mission_id)
        valid_interps = [i for i in interpretations if i.status in ("validated", "corrected")]

        # Build maturity profile
        profile: dict[str, list] = defaultdict(list)
        for interp in valid_interps:
            cap_id = interp.corrected_capacity_id or interp.capacity_id or "unknown"
            score = interp.corrected_maturity_score if interp.status == "corrected" else interp.maturity_score
            direction = interp.corrected_direction if interp.status == "corrected" else interp.direction
            profile[cap_id].append({"score": score, "direction": direction})

        # Load extracts for annex
        all_extracts = []
        verbatims = crud.list_verbatims_for_mission(session, mission_id)
        for verbatim in verbatims:
            extracts = crud.list_extracts(session, verbatim.verbatim_id)
            for extract in extracts:
                interps = crud.list_interpretations(session, extract.extract_id)
                for interp in interps:
                    all_extracts.append({
                        "text": extract.raw_text,
                        "capacity_id": interp.corrected_capacity_id or interp.capacity_id or "—",
                        "direction": interp.corrected_direction or interp.direction or "—",
                        "status": interp.status,
                    })

    # ── Cover page ────────────────────────────────────────────────────────────
    doc.add_heading(labels["cover_title"], level=0)
    cover_table = doc.add_table(rows=2, cols=2)
    cover_table.rows[0].cells[0].text = labels["client"]
    cover_table.rows[0].cells[1].text = client_name
    cover_table.rows[1].cells[0].text = labels["date"]
    cover_table.rows[1].cells[1].text = str(date.today())
    doc.add_page_break()

    # ── Maturity profile table ────────────────────────────────────────────────
    doc.add_heading(labels["profile_title"], level=1)
    if profile:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = labels["capacity"]
        hdr[1].text = labels["avg_score"]
        hdr[2].text = labels["nb_obs"]
        hdr[3].text = labels["directions"]

        for cap_id in sorted(profile.keys()):
            entries = profile[cap_id]
            scores = [e["score"] for e in entries if e["score"] is not None]
            avg = round(sum(scores) / len(scores), 1) if scores else "—"
            dirs = ", ".join(e["direction"] for e in entries if e["direction"])
            row = table.add_row().cells
            row[0].text = cap_id
            row[1].text = str(avg)
            row[2].text = str(len(entries))
            row[3].text = dirs
    doc.add_page_break()

    # ── Narrative ─────────────────────────────────────────────────────────────
    doc.add_heading(labels["narrative_title"], level=1)
    if report_content:
        for line in report_content.splitlines():
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.strip():
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("—")
    doc.add_page_break()

    # ── Annex: extracts ───────────────────────────────────────────────────────
    doc.add_heading(labels["extracts_title"], level=1)
    if all_extracts:
        annex_table = doc.add_table(rows=1, cols=4)
        annex_table.style = "Table Grid"
        ahdr = annex_table.rows[0].cells
        ahdr[0].text = labels["extract"]
        ahdr[1].text = labels["capacity"]
        ahdr[2].text = labels["directions"]
        ahdr[3].text = labels["validation"]
        for item in all_extracts:
            row = annex_table.add_row().cells
            row[0].text = item["text"][:200]
            row[1].text = item["capacity_id"]
            row[2].text = item["direction"]
            row[3].text = item["status"]

    output_path = Path(output_path)
    doc.save(output_path)
    return output_path
